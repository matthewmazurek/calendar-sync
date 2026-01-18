"""Word document reader for calendar files."""

import logging
import os
import re
import shutil
import subprocess
import tempfile
from datetime import date, datetime
from pathlib import Path
from typing import List

from docx import Document

from app.exceptions import IngestionError, InvalidYearError
from app.ingestion.summary import build_ingestion_summary
from app.models.event import Event
from app.models.ingestion import IngestionResult, RawIngestion
from app.models.template import CalendarTemplate, EventTypeConfig

logger = logging.getLogger(__name__)


class TypeMatcher:
    """Matches event text to event types using template configuration."""

    def __init__(self, template: CalendarTemplate):
        """Initialize matcher with template."""
        self.template = template
        # Build ordered list of type configs (order matters for matching)
        self.type_configs: list[tuple[str, EventTypeConfig]] = list(
            template.types.items()
        )

    def match_type(self, text: str) -> tuple[str | None, str | None]:
        """
        Match text to an event type and extract label.

        Returns:
            Tuple of (type_name, label) or (None, None) if no match
        """
        text_lower = text.lower()

        for type_name, config in self.type_configs:
            # Handle match patterns
            matches = config.match if isinstance(config.match, list) else [config.match]

            for pattern in matches:
                if config.match_mode == "regex":
                    match = re.search(pattern, text_lower, re.IGNORECASE)
                    if match:
                        label = None
                        if config.label:
                            # Extract label using label regex
                            label_match = re.search(config.label, text, re.IGNORECASE)
                            if label_match and label_match.lastindex >= 1:
                                label = label_match.group(1).strip()
                        return (type_name, label)
                else:  # contains mode
                    if pattern.lower() in text_lower:
                        label = None
                        if config.label:
                            # Extract label using label regex
                            label_match = re.search(config.label, text, re.IGNORECASE)
                            if label_match and label_match.lastindex >= 1:
                                label = label_match.group(1).strip()
                        return (type_name, label)

        return (None, None)

    def should_suppress(self, type_name: str | None) -> bool:
        """Check if the given type should be suppressed during ingestion."""
        if type_name is None:
            return False
        type_config = self.template.types.get(type_name)
        if type_config is None:
            return False
        return type_config.suppress

    def get_busy(self, type_name: str | None) -> bool:
        """Get the busy flag for the given type (default True)."""
        if type_name is None:
            return True
        type_config = self.template.types.get(type_name)
        if type_config is None:
            return True
        return type_config.busy

    def resolve_time_periods(
        self, text: str, type_name: str | None
    ) -> tuple[str | None, str | None]:
        """
        Resolve time periods (AM/PM) to actual times if applicable.

        Returns:
            Tuple of (start_time, end_time) or (None, None)
        """
        if type_name is None:
            return (None, None)

        type_config = self.template.types.get(type_name)
        if not type_config:
            return (None, None)

        defaults = self.template.defaults
        # Use type-specific time_periods if available, otherwise fall back to defaults
        time_periods = type_config.time_periods or defaults.time_periods

        # Check if text contains a time period indicator
        text_upper = text.upper()
        for period, (start, end) in time_periods.items():
            if period in text_upper:
                return (start, end)

        return (None, None)


# Map month names (uppercase) to month numbers
MONTH_MAP = {
    "JANUARY": 1,
    "FEBRUARY": 2,
    "MARCH": 3,
    "APRIL": 4,
    "MAY": 5,
    "JUNE": 6,
    "JULY": 7,
    "AUGUST": 8,
    "SEPTEMBER": 9,
    "OCTOBER": 10,
    "NOVEMBER": 11,
    "DECEMBER": 12,
}


def extract_year_from_header(header: str) -> int | None:
    """Extract the year from a month header (e.g., "JANUARY 2025")."""
    parts = header.split()
    if len(parts) < 2 or parts[0] not in MONTH_MAP:
        return None
    try:
        return int(parts[1])
    except ValueError:
        return None


def extract_revised_date(doc: Document) -> date | None:
    """Extract revised date from document header.

    Looks for pattern like "Revised December 16, 2025" in headers or paragraphs.
    """
    # Pattern: "Revised December 16, 2025" (matches format in header)
    pattern = r"Revised\s+([A-Za-z]+)\s+(\d+),?\s+(\d{4})"

    # Check headers first (most common location)
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                text = para.text
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    month_name = match.group(1).upper()
                    day = int(match.group(2))
                    year = int(match.group(3))
                    if month_name in MONTH_MAP:
                        try:
                            revised_date = datetime(
                                year, MONTH_MAP[month_name], day
                            ).date()
                            logger.info(
                                f"Extracted revised date from header: {revised_date}"
                            )
                            return revised_date
                        except ValueError:
                            pass

    logger.debug("No revised date found in document")
    return None


def extract_time_ranges(text: str) -> list[tuple[str, str, str]] | None:
    """
    Extract multiple time ranges from text like "CCSC 0730-1200 and 1230-1630".
    
    Returns:
        List of (start, end, connector) tuples if multiple ranges found, None otherwise.
        Connector is the text between ranges (e.g., "and", "&", "+").
    """
    # Pattern: Find all HHMM-HHMM time ranges
    time_range_pattern = r'(\d{4})-(\d{4})'
    matches = list(re.finditer(time_range_pattern, text))
    
    if len(matches) < 2:
        return None
    
    # Extract the ranges and check for conjunctions between them
    ranges = []
    for i, match in enumerate(matches):
        start = match.group(1)
        end = match.group(2)
        
        # Check if there's a conjunction before the next range
        if i < len(matches) - 1:
            between_start = match.end()
            between_end = matches[i + 1].start()
            between_text = text[between_start:between_end].strip()
            
            # Check for conjunction patterns (word boundaries don't work with & and +)
            if re.search(r'(\band\b|&|\+)', between_text, re.IGNORECASE):
                ranges.append((start, end, between_text))
            else:
                # No conjunction found, not a multi-range event
                return None
        else:
            # Last range
            ranges.append((start, end, ""))
    
    return ranges if len(ranges) >= 2 else None


def extract_time_periods(text: str) -> list[str] | None:
    """
    Extract multiple time periods (AM/PM) from text like "CCSC AM and PM".
    
    Returns:
        List of period strings if multiple periods found with conjunctions, None otherwise.
    """
    # Find AM/PM occurrences
    period_pattern = r'\b(AM|PM)\b'
    matches = list(re.finditer(period_pattern, text, re.IGNORECASE))
    
    if len(matches) < 2:
        return None
    
    # Check if there's a conjunction between periods
    periods = []
    for i, match in enumerate(matches):
        period = match.group(1).upper()
        
        if i < len(matches) - 1:
            between_start = match.end()
            between_end = matches[i + 1].start()
            between_text = text[between_start:between_end].strip()
            
            # Check for conjunction patterns (word boundaries don't work with & and +)
            if re.search(r'(\band\b|&|\+)', between_text, re.IGNORECASE):
                periods.append(period)
            else:
                # No conjunction found
                return None
        else:
            # Last period
            periods.append(period)
    
    return periods if len(periods) >= 2 else None


def parse_cell_events(
    cell_text: str, month: int, year: int, type_matcher: TypeMatcher | None = None
) -> List[dict]:
    """Parse events from a calendar cell."""
    events = []
    lines = [ln.strip() for ln in cell_text.split("\n") if ln.strip()]
    if not lines:
        return events

    # First line may be "day [first event]" or just "day"
    m_day = re.match(r"^(\d+)\s*(.*)$", lines[0])
    if not m_day:
        return events
    day = int(m_day.group(1))
    first_ev = m_day.group(2).strip()

    # Special case: If we're in December and this is New Year's Day, it belongs to January
    if month == 12 and day == 1 and first_ev.lower().startswith("new year"):
        date_str = f"{year + 1}-01-01"
    else:
        date_str = f"{year}-{month:02d}-{day:02d}"

    # Collect all event lines
    ev_lines = []
    if first_ev:
        ev_lines.append(first_ev)
    ev_lines.extend(lines[1:])

    # For each event line, split on commas and parse times
    for ln in ev_lines:
        for part in ln.split(","):
            ev = part.strip()
            if not ev:
                continue
            
            # Check for multiple time ranges first (e.g., "CCSC 0730-1200 and 1230-1630")
            multi_ranges = extract_time_ranges(ev)
            if multi_ranges:
                # Extract title (everything before the first time range)
                m_title = re.match(r"(.+?)\s+\d{4}-\d{4}", ev)
                title = m_title.group(1).strip() if m_title else ev
                
                # Create separate events for each time range
                for start, end, _ in multi_ranges:
                    event_dict = {
                        "title": title,
                        "start": start,
                        "end": end,
                        "date": date_str,
                    }
                    
                    # Apply template matching if available
                    if type_matcher:
                        type_name, label = type_matcher.match_type(title)
                        if type_name:
                            event_dict["type"] = type_name
                            event_dict["busy"] = type_matcher.get_busy(type_name)
                        if label:
                            event_dict["label"] = label
                    
                    events.append(event_dict)
                continue
            
            # If there's a time range like "Endo 1230-1630" or "Clinic 1230-1630 with Carmen"
            m_time = re.match(r"(.+?)\s+(\d{4})-(\d{4})(?:\s+(.+))?", ev)
            if m_time:
                title = m_time.group(1).strip()
                start = m_time.group(2)
                end = m_time.group(3)
                additional_text = m_time.group(4) if m_time.group(4) else ""

                full_title = title
                if additional_text:
                    full_title = f"{title} ({additional_text.strip()})"

                event_dict = {
                    "title": full_title,
                    "start": start,
                    "end": end,
                    "date": date_str,
                }

                # Apply template matching if available
                if type_matcher:
                    type_name, label = type_matcher.match_type(full_title)
                    if type_name:
                        # Use template type name directly (user-defined, no enum conversion)
                        event_dict["type"] = type_name
                        event_dict["busy"] = type_matcher.get_busy(type_name)
                    if label:
                        event_dict["label"] = label

                    # Check for time period resolution (e.g., "CCSC AM")
                    period_start, period_end = type_matcher.resolve_time_periods(
                        full_title, type_name
                    )
                    if period_start and period_end:
                        # Override times with period times
                        event_dict["start"] = period_start
                        event_dict["end"] = period_end

                events.append(event_dict)
            else:
                # Check for multiple time periods (e.g., "CCSC AM and PM")
                multi_periods = extract_time_periods(ev)
                if multi_periods and type_matcher:
                    # Extract base title (remove AM/PM indicators)
                    title = re.sub(r'\s+(AM|PM)\s+(and|&|\+)\s+(AM|PM).*', '', ev, flags=re.IGNORECASE).strip()
                    title = re.sub(r'\s+(AM|PM)$', '', title, flags=re.IGNORECASE).strip()
                    
                    # Apply template matching
                    type_name, label = type_matcher.match_type(ev)
                    
                    # Create separate events for each period
                    for period in multi_periods:
                        period_start, period_end = type_matcher.resolve_time_periods(
                            period, type_name
                        )
                        
                        if period_start and period_end:
                            event_dict = {
                                "title": title,
                                "start": period_start,
                                "end": period_end,
                                "date": date_str,
                            }
                            if type_name:
                                event_dict["type"] = type_name
                                event_dict["busy"] = type_matcher.get_busy(type_name)
                            if label:
                                event_dict["label"] = label
                            events.append(event_dict)
                    continue
                
                # All-day or untimed event
                event_dict = {"title": ev, "date": date_str}

                # Apply template matching if available
                if type_matcher:
                    type_name, label = type_matcher.match_type(ev)
                    if type_name:
                        # Use template type name directly (user-defined, no enum conversion)
                        event_dict["type"] = type_name
                        event_dict["busy"] = type_matcher.get_busy(type_name)
                    if label:
                        event_dict["label"] = label

                    # Check for time period resolution (e.g., "CCSC AM")
                    period_start, period_end = type_matcher.resolve_time_periods(
                        ev, type_name
                    )
                    if period_start and period_end:
                        # Convert to timed event with period times
                        event_dict["start"] = period_start
                        event_dict["end"] = period_end

                events.append(event_dict)

    return events


def normalize_to_docx(path: str | Path) -> str:
    """Convert .doc to .docx if needed."""
    path = str(path)
    root, ext = os.path.splitext(path)
    if ext.lower() == ".doc":
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_doc = os.path.join(temp_dir, os.path.basename(path))
            shutil.copy2(path, temp_doc)

            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    temp_dir,
                    temp_doc,
                ],
                check=True,
            )

            temp_docx = os.path.splitext(temp_doc)[0] + ".docx"

            with tempfile.NamedTemporaryFile(
                suffix=".docx", delete=False
            ) as final_docx:
                shutil.copy2(temp_docx, final_docx.name)
                return final_docx.name
    elif ext.lower() == ".docx":
        return path
    else:
        raise ValueError(f"Unsupported extension: {ext}")


class WordReader:
    """Reader for Word document calendar files."""

    def read(
        self, path: Path, template: CalendarTemplate | None = None
    ) -> IngestionResult:
        """Read calendar from Word document."""
        try:
            docx_path = normalize_to_docx(path)
            try:
                raw = self._read_docx(docx_path, template)
                return IngestionResult(
                    raw=raw, summary=build_ingestion_summary(raw)
                )
            finally:
                # Clean up temporary docx file if it was created
                if docx_path != str(path):
                    try:
                        os.unlink(docx_path)
                    except OSError:
                        pass
        except Exception as e:
            raise IngestionError(f"Failed to read Word document: {e}") from e

    def _read_docx(self, docx_path: str, template: CalendarTemplate | None) -> RawIngestion:
        """Read calendar from .docx file."""
        logger.info(f"Reading Word document: {docx_path}")
        if template:
            logger.info(f"Using template: {template.name} (version {template.version})")
        else:
            logger.info("No template provided")
        doc = Document(str(docx_path))
        events = []

        # Create type matcher if template provided
        type_matcher = TypeMatcher(template) if template else None

        # Extract revised date
        revised_date = extract_revised_date(doc)
        if revised_date:
            logger.info(f"Extracted revised date from Word document: {revised_date}")
        else:
            logger.debug("No revised date found in Word document")

        # Single big table: 12 months × 8 rows per month
        if not doc.tables:
            raise IngestionError("Document contains no tables")

        table = doc.tables[0]

        # Extract year from the first month's header
        first_header = table.rows[0].cells[0].text.strip().upper()
        year = extract_year_from_header(first_header)
        if year is None:
            raise IngestionError("Could not determine year from document")

        logger.info(f"Extracted year: {year}")

        cutoff_date = datetime(year, 12, 31)

        for month_idx in range(12):
            start_row = month_idx * 8
            if start_row >= len(table.rows):
                continue

            header = table.rows[start_row].cells[0].text.strip().upper()
            parts = header.split()
            if len(parts) < 2 or parts[0] not in MONTH_MAP or parts[1] != str(year):
                continue

            month = MONTH_MAP[parts[0]]

            # Iterate the 6 week rows
            for row in table.rows[start_row + 2 : start_row + 8]:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        events.extend(
                            parse_cell_events(cell_text, month, year, type_matcher)
                        )

        # Filter out any events after December 31 of the target year
        events = [
            e for e in events if datetime.strptime(e["date"], "%Y-%m-%d") <= cutoff_date
        ]

        # Remove misattributed New Year's Day on Dec 1
        events = [
            e
            for e in events
            if not (
                e["title"].lower().startswith("new year")
                and e["date"] == f"{year}-12-01"
            )
        ]

        # Convert to Pydantic Event models, filtering suppressed types
        event_models = []
        suppressed_count = 0
        for event_dict in events:
            # Check if this event type should be suppressed
            if type_matcher and type_matcher.should_suppress(event_dict.get("type")):
                suppressed_count += 1
                logger.debug(f"Suppressing event: {event_dict.get('title')}")
                continue
            try:
                event_models.append(Event(**event_dict))
            except Exception as e:
                raise IngestionError(
                    f"Failed to create event from {event_dict}: {e}"
                ) from e

        if suppressed_count > 0:
            logger.info(f"Suppressed {suppressed_count} events based on template rules")
        logger.info(f"Created {len(event_models)} event models from Word document")

        # Validate single year (Word documents are expected to be single-year schedules)
        years = {event.date.year for event in event_models}
        if len(years) > 1:
            raise InvalidYearError(
                f"Word document contains events from multiple years: {years}. "
                "Word documents must contain events from a single year."
            )

        return RawIngestion(events=event_models, revised_at=revised_date)
